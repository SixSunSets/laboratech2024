import pandas as pd
import pandas as pd
from docx import Document
import json
from docx2pdf import convert
from docx.shared import Pt 
import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

df = pd.read_csv('data/data.csv')

depa= pd.read_excel('data/diccionario.xlsx',sheet_name='Departamento',skiprows=3)
depa = depa[['Código','DEPARTAMENTO']].loc[:24,:]
depa['Código'] = depa['Código'].astype(int)


sector= pd.read_excel('data/diccionario.xlsx',sheet_name='Sector',skiprows=2)
sector.columns = sector.iloc[0]
sector.drop(index=0, inplace=True)
sector.reset_index(drop=True,inplace=True)
sector['Código'] =sector['Código'].astype(int)


def replace_text(doc, replacements):
    def replace_in_paragraph(paragraph, replacements):
        full_text = ''.join(run.text for run in paragraph.runs)
        
        for old_text, new_text in replacements.items():
            if old_text in full_text:
                full_text = full_text.replace(old_text, new_text)
        
        for run in paragraph.runs:
            run.clear()

        paragraph.add_run(full_text)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)

def convertir_fecha(valor):
    valor_str = str(valor)
    año = valor_str[:4]
    mes_num = valor_str[4:6]
    
    meses = {
        '01': 'Enero', '02': 'Febrero', '03': 'Marzo', '04': 'Abril',
        '05': 'Mayo', '06': 'Junio', '07': 'Julio', '08': 'Agosto',
        '09': 'Septiembre', '10': 'Octubre', '11': 'Noviembre', '12': 'Diciembre'
    }
    
    mes_nombre = meses.get(mes_num, 'Mes inválido')
    
    return f"{año} - {mes_nombre}"

def return_sector(id):
    try:
        return f"{sector[sector['Código']==id]['Descripción'].values[0]}"
    except:
        return '-'

def return_depa(id):
    try:
        return f"{depa[depa['Código']==id]['DEPARTAMENTO'].values[0]}"
    except:
        return '-'

def apply_format_to_cell(cell, bold=False, font_size=None):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if bold:
                run.bold = True
            if font_size:
                run.font.size = Pt(font_size)

def apply_format_to_text(doc, text_to_format, bold=False, font_size=None):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if text_to_format in run.text:
                if bold:
                    run.bold = True
                if font_size:
                    run.font.size = Pt(font_size)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if text_to_format in run.text:
                            if bold:
                                run.bold = True
                            if font_size:
                                run.font.size = Pt(font_size)



# ----------------------------------------------------------------------

# 202401 - 202406

fechas = {"fecha_inicio":202401,'fecha_fin':202406}

data = df[(df['V_NUMPERIOD']>=fechas['fecha_inicio']) & (df['V_NUMPERIOD']<=fechas['fecha_fin'])]

data.loc[:,'V_NUMPERIOD'] = data['V_NUMPERIOD'].astype(str)

data.loc[:,'V_NUMPERIOD'] = data['V_NUMPERIOD'].apply(lambda x: f"{x[:4]}-{x[4:]}")


data_tc = data[['V_NUMPERIOD','ntrab_tc_indet','ntrab_tc_nattemp','ntrab_tc_natacc','ntrab_tc_obrserv','ntrab_tc_tiempar','ntrab_tc_otros','ntrab_tc_noprec']]
data_tc.columns = ["periodo","A plazo indeterminado","De naturaleza temporal",'De naturaleza accidental','Contrato de obra','Contrato a tiempo parcial','Otros tipos de contrato','No precisa' ]
data_tc = data_tc.sort_values('periodo').set_index('periodo').T
data_tc = data_tc.reset_index()
data_tc.rename(columns={'index':'Tipo de Contrato'},inplace=True)


data_rl = data[['V_NUMPERIOD','ntrab_rl_privgen','ntrab_rl_micro','ntrab_rl_peq','ntrab_rl_agrar','ntrab_rl_otros','ntrab_rl_noprec']]
data_rl.columns = ["periodo","General",'Micro Empresa',"Pequeña Empresa",'Agrario','Otros','No precisa']
data_rl = data_rl.sort_values('periodo').set_index('periodo').T
data_rl = data_rl.reset_index()
data_rl.rename(columns={'index':'Régimen Laboral'},inplace=True)


data_rp = data[['V_NUMPERIOD','ntrab_rp_spp','ntrab_rp_snp','ntrab_rp_otros','ntrab_rp_noprecrp']]
data_rp.columns = ["periodo","Sist. Priv. de Pensiones",'Sist. Nac. de Pensiones',"Otros régim. pensionarios",'No precisa']
data_rp = data_rp.sort_values('periodo').set_index('periodo').T
data_rp = data_rp.reset_index()
data_rp.rename(columns={'index':'Régimen Pensionario'},inplace=True)


data_afil = data[['V_NUMPERIOD','ntrab_afil_essalud','ntrab_afil_essalud_agrar','ntrab_afil_sis','ntrab_afil_noprec']]
data_afil.columns = ["periodo","ESSALUD",'ESSALUD Agrario',"SIS",'No precisa']
data_afil = data_afil.sort_values('periodo').set_index('periodo').T
data_afil = data_afil.reset_index()
data_afil.rename(columns={'index':'Régimen de Salud'},inplace=True)


data_ran = data[['V_NUMPERIOD','ntrab_18mas','ntrab_menos18','ntrab_edad_noprec']]
data_ran.columns = ["periodo","De 18 a más",'Menores de 18 años','No precisa']
data_ran = data_ran.sort_values('periodo').set_index('periodo').T
data_ran = data_ran.reset_index()
data_ran.rename(columns={'index':'Rango de Edades'},inplace=True)


data_gen = data[['V_NUMPERIOD','ntrab_hombres','ntrab_mujeres','ntrab_sexo_noprec']]
data_gen.columns = ["periodo","Femenino",'Masculino','No precisa']
data_gen = data_gen.sort_values('periodo').set_index('periodo').T
data_gen = data_gen.reset_index()
data_gen.rename(columns={'index':'Género'},inplace=True)


data_sind = data[['V_NUMPERIOD','ntrab_nosind','ntrab_sind']]
data_sind.columns = ["periodo","Sindicalizados",'No sindicalizados']
data_sind = data_sind.sort_values('periodo').set_index('periodo').T
data_sind = data_sind.reset_index()
data_sind.rename(columns={'index':'Afiliación Sindical'},inplace=True)



data_sctr = data[['V_NUMPERIOD','ntrab_consctr','ntrab_sinsctr','ntrab_sctr_noprec']]
data_sctr.columns = ["periodo","Con SCTR",'Sin SCTR','No precisa']
data_sctr = data_sctr.sort_values('periodo').set_index('periodo').T
data_sctr = data_sctr.reset_index()
data_sctr.rename(columns={'index':'Afiliación al SCTR'},inplace=True)





template_path = "template_month.docx"
doc = Document(template_path)

fila = df[df['V_NUMPERIOD'] == fechas['fecha_inicio']]

# Obtención de valores para reemplazo
Z0, Z8, Z1, Z2, Z3, Z4, Z5, Z7 = (
    convertir_fecha(fechas['fecha_inicio']),
    convertir_fecha(fechas['fecha_fin']),
    fila['ID_EMPRESA'].values[0],
    return_depa(fila['Departamento_ct'].values[0]),
    return_sector(fila['sector'].values[0]),
    fila['anio_anti'].values[0],
    fila['numtra'].values[0],
    fila['sst'].values[0]
)

# Diccionario de reemplazos
replacements = {
    'Z0': str(Z0),
    'Z1': str(Z1),
    'Z2': str(Z2),
    'Z3': str(Z3),
    'Z4': str(Z4),
    'Z5': str(Z5),
    'Z7': str(Z7),
    'Z8': str(Z8)
}

for paragraph in doc.paragraphs:
    for key, value in replacements.items():
        if key in paragraph.text:
            paragraph.text = paragraph.text.replace(key, value)

# Crear la primera tabla usando data_tc
tabla_tc = doc.add_table(rows=1 + len(data_tc), cols=len(data_tc.columns))
tabla_tc.style = 'Table Grid'

for j, col_name in enumerate(data_tc.columns):
    cell = tabla_tc.rows[0].cells[j]
    cell.text = str(col_name)
    # Centramos el texto del encabezado
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

for i, row in data_tc.iterrows():
    for j, value in enumerate(row):
        cell = tabla_tc.rows[i + 1].cells[j]
        cell.text = str(value)  # Convertir todos los valores a cadena
        # Centramos el texto de las celdas
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

for paragraph in doc.paragraphs:
    if "{{Tabla}}" in paragraph.text:
        table_location = paragraph
        table_location.text = ""
        table_location._element.addnext(tabla_tc._tbl)  # Inserta la tabla en lugar del marcador

# Crear un párrafo vacío como espacio entre las tablas
paragraph = doc.add_paragraph()
paragraph.add_run("")  # Espacio vacío
paragraph_format = paragraph.paragraph_format
paragraph_format.space_after = Pt(8)  # Ajusta el espacio después del párrafo en puntos (opcional)



# Crear la segunda tabla usando data_rl y añadirla después de la primera tabla
tabla_rl = doc.add_table(rows=1 + len(data_rl), cols=len(data_rl.columns))
tabla_rl.style = 'Table Grid'

# Encabezados de la segunda tabla
for j, col_name in enumerate(data_rl.columns):
    cell = tabla_rl.rows[0].cells[j]
    cell.text = str(col_name)
    # Centramos el texto del encabezado
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Llenar las filas de la segunda tabla con datos
for i, row in data_rl.iterrows():
    for j, value in enumerate(row):
        cell = tabla_rl.rows[i + 1].cells[j]
        cell.text = str(value)  # Convertir todos los valores a cadena
        # Centramos el texto de las celdas
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Insertar la segunda tabla justo después de la primera
tabla_tc._tbl.addnext(tabla_rl._tbl)

paragraph._element.addnext(tabla_rl._tbl)


# Crear la segunda tabla usando data_rl y añadirla después de la primera tabla
tabla_rp = doc.add_table(rows=1 + len(data_rp), cols=len(data_rp.columns))
tabla_rp.style = 'Table Grid'

# Encabezados de la tercera tabla
for j, col_name in enumerate(data_rp.columns):
    cell = tabla_rp.rows[0].cells[j]
    cell.text = str(col_name)
    # Centramos el texto del encabezado
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Llenar las filas de la segunda tabla con datos
for i, row in data_rp.iterrows():
    for j, value in enumerate(row):
        cell = tabla_rp.rows[i + 1].cells[j]
        cell.text = str(value)  # Convertir todos los valores a cadena
        # Centramos el texto de las celdas
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Insertar la segunda tabla justo después de la primera
tabla_rl._tbl.addnext(tabla_rp._tbl)


# Crear un párrafo vacío como espacio entre las tablas
paragraph = doc.add_paragraph()
paragraph.add_run("")  # Espacio vacío
paragraph_format = paragraph.paragraph_format
paragraph_format.space_after = Pt(8)  # Ajusta el espacio después del párrafo en puntos (opcional)

paragraph._element.addnext(tabla_rp._tbl)

# Crear la segunda tabla usando data_rl y añadirla después de la primera tabla
tabla_afil = doc.add_table(rows=1 + len(data_afil), cols=len(data_afil.columns))
tabla_afil.style = 'Table Grid'

# Encabezados de la tercera tabla
for j, col_name in enumerate(data_afil.columns):
    cell = tabla_afil.rows[0].cells[j]
    cell.text = str(col_name)
    # Centramos el texto del encabezado
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Llenar las filas de la segunda tabla con datos
for i, row in data_afil.iterrows():
    for j, value in enumerate(row):
        cell = tabla_afil.rows[i + 1].cells[j]
        cell.text = str(value)  # Convertir todos los valores a cadena
        # Centramos el texto de las celdas
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Insertar la segunda tabla justo después de la primera
tabla_rp._tbl.addnext(tabla_afil._tbl)

# Crear un párrafo vacío como espacio entre las tablas
paragraph = doc.add_paragraph()
paragraph.add_run("")  # Espacio vacío
paragraph_format = paragraph.paragraph_format
paragraph_format.space_after = Pt(8)  # Ajusta el espacio después del párrafo en puntos (opcional)

paragraph._element.addnext(tabla_afil._tbl)


# Crear la segunda tabla usando data_rl y añadirla después de la primera tabla
tabla_ran = doc.add_table(rows=1 + len(data_ran), cols=len(data_ran.columns))
tabla_ran.style = 'Table Grid'

# Encabezados de la tercera tabla
for j, col_name in enumerate(data_ran.columns):
    cell = tabla_ran.rows[0].cells[j]
    cell.text = str(col_name)
    # Centramos el texto del encabezado
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Llenar las filas de la segunda tabla con datos
for i, row in data_ran.iterrows():
    for j, value in enumerate(row):
        cell = tabla_ran.rows[i + 1].cells[j]
        cell.text = str(value)  # Convertir todos los valores a cadena
        # Centramos el texto de las celdas
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Insertar la segunda tabla justo después de la primera
tabla_afil._tbl.addnext(tabla_ran._tbl)

# Crear un párrafo vacío como espacio entre las tablas
paragraph = doc.add_paragraph()
paragraph.add_run("")  # Espacio vacío
paragraph_format = paragraph.paragraph_format
paragraph_format.space_after = Pt(8)  # Ajusta el espacio después del párrafo en puntos (opcional)

paragraph._element.addnext(tabla_ran._tbl)

# Crear la segunda tabla usando data_rl y añadirla después de la primera tabla
tabla_gen = doc.add_table(rows=1 + len(data_gen), cols=len(data_gen.columns))
tabla_gen.style = 'Table Grid'

# Encabezados de la tercera tabla
for j, col_name in enumerate(data_gen.columns):
    cell = tabla_gen.rows[0].cells[j]
    cell.text = str(col_name)
    # Centramos el texto del encabezado
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Llenar las filas de la segunda tabla con datos
for i, row in data_gen.iterrows():
    for j, value in enumerate(row):
        cell = tabla_gen.rows[i + 1].cells[j]
        cell.text = str(value)  # Convertir todos los valores a cadena
        # Centramos el texto de las celdas
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Insertar la segunda tabla justo después de la primera
tabla_ran._tbl.addnext(tabla_gen._tbl)

# Crear un párrafo vacío como espacio entre las tablas
paragraph = doc.add_paragraph()
paragraph.add_run("")  # Espacio vacío
paragraph_format = paragraph.paragraph_format
paragraph_format.space_after = Pt(8)  # Ajusta el espacio después del párrafo en puntos (opcional)

paragraph._element.addnext(tabla_gen._tbl)


# Crear la segunda tabla usando data_rl y añadirla después de la primera tabla
tabla_sind = doc.add_table(rows=1 + len(data_sind), cols=len(data_sind.columns))
tabla_sind.style = 'Table Grid'

# Encabezados de la tercera tabla
for j, col_name in enumerate(data_sind.columns):
    cell = tabla_sind.rows[0].cells[j]
    cell.text = str(col_name)
    # Centramos el texto del encabezado
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Llenar las filas de la segunda tabla con datos
for i, row in data_sind.iterrows():
    for j, value in enumerate(row):
        cell = tabla_sind.rows[i + 1].cells[j]
        cell.text = str(value)  # Convertir todos los valores a cadena
        # Centramos el texto de las celdas
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Insertar la segunda tabla justo después de la primera
tabla_gen._tbl.addnext(tabla_sind._tbl)



# Crear un párrafo vacío como espacio entre las tablas
paragraph = doc.add_paragraph()
paragraph.add_run("")  # Espacio vacío
paragraph_format = paragraph.paragraph_format
paragraph_format.space_after = Pt(8)  # Ajusta el espacio después del párrafo en puntos (opcional)

paragraph._element.addnext(tabla_sind._tbl)

# Crear la segunda tabla usando data_rl y añadirla después de la primera tabla
tabla_sctr = doc.add_table(rows=1 + len(data_sctr), cols=len(data_sctr.columns))
tabla_sctr.style = 'Table Grid'

# Encabezados de la tercera tabla
for j, col_name in enumerate(data_sctr.columns):
    cell = tabla_sctr.rows[0].cells[j]
    cell.text = str(col_name)
    # Centramos el texto del encabezado
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Llenar las filas de la segunda tabla con datos
for i, row in data_sctr.iterrows():
    for j, value in enumerate(row):
        cell = tabla_sctr.rows[i + 1].cells[j]
        cell.text = str(value)  # Convertir todos los valores a cadena
        # Centramos el texto de las celdas
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Insertar la segunda tabla justo después de la primera
tabla_sind._tbl.addnext(tabla_sctr._tbl)


# Crear un párrafo vacío como espacio entre las tablas
paragraph = doc.add_paragraph()
paragraph.add_run("")  # Espacio vacío
paragraph_format = paragraph.paragraph_format
paragraph_format.space_after = Pt(8)  # Ajusta el espacio después del párrafo en puntos (opcional)

paragraph._element.addnext(tabla_sctr._tbl)


# Guardar el documento modificado
output_path = f"planilla_Electrónica_meses_{str(fechas['fecha_inicio'])}-{str(fechas['fecha_fin'])}.docx"
doc.save(output_path)

print(f"Documento guardado en: {output_path}")

formatos = [
("Planilla Electrónica", {'bold': False, 'font_size': 18}),
("Empresa:", {'bold': True, 'font_size': 12}),
(f"Departamento", {'bold': True, 'font_size': 12}),
("Sector", {'bold': True, 'font_size': 12}),
("Años en el mercado", {'bold': True, 'font_size': 12}),
("Nº de trabajadores", {'bold': True, 'font_size': 12}),
("SST", {'bold': True, 'font_size': 12}),

("Tipo de Contrato", {'bold': True, 'font_size': 11}),
("A plazo indeterminado", {'bold': True, 'font_size': 11}),
("De naturaleza temporal", {'bold': True, 'font_size': 11}),
('De naturaleza accidental', {'bold': True, 'font_size': 11}),
('Contrato de obra', {'bold': True, 'font_size': 11}),
('Contrato a tiempo parcial', {'bold': True, 'font_size': 11}),
('Otros tipos de contrato', {'bold': True, 'font_size': 11}),
('No precisa', {'bold': True, 'font_size': 11}),

("Régimen Laboral", {'bold': True, 'font_size': 11}),
("General", {'bold': True, 'font_size': 11}),
("Micro Empresa", {'bold': True, 'font_size': 11}),
('Pequeña Empresa', {'bold': True, 'font_size': 11}),
('Agrario', {'bold': True, 'font_size': 11}),
('Otros', {'bold': True, 'font_size': 11}),

("Régimen Pensionario", {'bold': True, 'font_size': 11}),
('Sist. Priv. de Pensiones', {'bold': True, 'font_size': 11}),
('Sist. Nac. de Pensiones', {'bold': True, 'font_size': 11}),
('Otros régim. pensionarios', {'bold': True, 'font_size': 11}),


("Régimen de Salud", {'bold': True, 'font_size': 11}),
('ESSALUD', {'bold': True, 'font_size': 11}),
('ESSALUD Agrario', {'bold': True, 'font_size': 11}),
('SIS', {'bold': True, 'font_size': 11}),

("Rango de Edades", {'bold': True, 'font_size': 11}),
('De 18 a más', {'bold': True, 'font_size': 11}),
('Menores de 18 años', {'bold': True, 'font_size': 11}),

("Género", {'bold': True, 'font_size': 11}),
('Femenino', {'bold': True, 'font_size': 11}),
('Masculino', {'bold': True, 'font_size': 11}),

("Afiliación Sindical", {'bold': True, 'font_size': 11}),
('Sindicalizados', {'bold': True, 'font_size': 11}),
('No sindicalizados', {'bold': True, 'font_size': 11}),

("Afiliación al SCTR", {'bold': True, 'font_size': 11}),
('Con SCTR', {'bold': True, 'font_size': 11}),
('Sin SCTR', {'bold': True, 'font_size': 11}),

('Multas Laborales', {'bold': True, 'font_size': 12}),
('Nº de Trabj. Afectados', {'bold': True, 'font_size': 11}),
('Monto', {'bold': True, 'font_size': 11}),
('Fecha Multa', {'bold': True, 'font_size': 11}),

(f'{data_tc.columns[1][:4]}-', {'bold': True, 'font_size': 11}),
]


doc.add_paragraph("")

titulo = doc.add_paragraph("Multas Laborales")

df_tabla = pd.read_csv('data/data2.csv')

df_tabla.loc[:,'V_NUMPERIOD'] = df_tabla['V_NUMPERIOD'].astype(str)
df_tabla.loc[:,'V_NUMPERIOD'] = df_tabla['V_NUMPERIOD'].apply(lambda x: f"{x[:4]}-{x[4:]}")

data_multas = df_tabla[['V_NUMPERIOD', 'TRAB_AFEC', 'MONTO_MULTA']]
data_multas.columns = ["periodo","Nº de Trabj. Afectados","Monto"]
data_multas = data_multas.sort_values('periodo').set_index('periodo').T
data_multas = data_multas.reset_index()
data_multas.rename(columns={'index':'Fecha Multa'},inplace=True)

tabla_multas = doc.add_table(rows=1 + len(data_multas), cols=len(data_multas.columns))
tabla_multas.style = 'Table Grid'

for j, col_name in enumerate(data_multas.columns):
    cell = tabla_multas.rows[0].cells[j]
    cell.text = str(col_name)
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

for i, row in data_multas.iterrows():
    for j, value in enumerate(row):
        cell = tabla_multas.rows[i + 1].cells[j]
        cell.text = str(value)  # Convertir todos los valores a cadena
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



for texto, propiedades in formatos:
        apply_format_to_text(doc, texto, **propiedades)


doc.save(f"planilla_Electrónica_meses_{str(fechas['fecha_inicio'])}-{str(fechas['fecha_fin'])}.docx")
nombre_docx = f"planilla_Electrónica_meses_{str(fechas['fecha_inicio'])}-{str(fechas['fecha_fin'])}.docx"
convert(nombre_docx)

# Nombre del archivo generado dinámicamente
file_name = f"planilla_Electrónica_meses_{str(fechas['fecha_inicio'])}-{str(fechas['fecha_fin'])}.docx"

# Verifica si el archivo existe antes de intentar eliminarlo
if os.path.exists(file_name):
    os.remove(file_name)
    print(f"Archivo eliminado: {file_name}")
else:
    print(f"El archivo no existe: {file_name}")