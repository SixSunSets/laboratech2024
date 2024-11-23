import pandas as pd
import pandas as pd
from docx import Document
import json
from docx2pdf import convert
from docx.shared import Pt 
import os
from docx.enum.text import WD_BREAK
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

df = pd.read_csv('C:/Users/alexa/Dataton/Labora-Tech-2024/backend/generate_pdf_planilla/data/data.csv')

depa= pd.read_excel('C:/Users/alexa/Dataton/Labora-Tech-2024/backend/generate_pdf_planilla/data/diccionario.xlsx',sheet_name='Departamento',skiprows=3)
depa = depa[['Código','DEPARTAMENTO']].loc[:24,:]
depa['Código'] = depa['Código'].astype(int)


sector= pd.read_excel('C:/Users/alexa/Dataton/Labora-Tech-2024/backend/generate_pdf_planilla/data/diccionario.xlsx',sheet_name='Sector',skiprows=2)
sector.columns = sector.iloc[0]
sector.drop(index=0, inplace=True)
sector.reset_index(drop=True,inplace=True)
sector['Código'] =sector['Código'].astype(int)

def replace_text(doc, replacements):
    def replace_in_paragraph(paragraph, replacements):
        full_text = ''.join(run.text for run in paragraph.runs)
        
        for old_text, new_text in replacements.items():
            if old_text in full_text:
                full_text = full_text.replace(old_text, new_text)
        
        for run in paragraph.runs:
            run.clear()

        paragraph.add_run(full_text)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)


def convertir_fecha(valor):
    valor_str = str(valor)
    año = valor_str[:4]
    mes_num = valor_str[4:6]
    
    meses = {
        '01': 'Enero', '02': 'Febrero', '03': 'Marzo', '04': 'Abril',
        '05': 'Mayo', '06': 'Junio', '07': 'Julio', '08': 'Agosto',
        '09': 'Septiembre', '10': 'Octubre', '11': 'Noviembre', '12': 'Diciembre'
    }
    
    mes_nombre = meses.get(mes_num, 'Mes inválido')
    
    return f"{año} - {mes_nombre}"

def return_sector(id):
    try:
        return f"{sector[sector['Código']==id]['Descripción'].values[0]}"
    except:
        return '-'

def return_depa(id):
    try:
        return f"{depa[depa['Código']==id]['DEPARTAMENTO'].values[0]}"
    except:
        return '-'

def process_text(row, fechas, doc):
    Z0,Z1,Z2,Z3,Z4,Z5,Z6,Z7 = convertir_fecha(row['V_NUMPERIOD'].values[0]), row['ID_EMPRESA'].values[0],return_depa(row['Departamento_ct'].values[0]), return_sector(row['sector'].values[0]), row['anio_anti'].values[0],row['numtra'].values[0], round(row['costosal'].values[0],2), row['sst'].values[0]
    
    # Variables A,B
    A0, A1, A2, A3, A4, A5 = row['ntrab_rl_privgen'].values[0], row['ntrab_rl_peq'].values[0], row['ntrab_rl_micro'].values[0], row['ntrab_rl_agrar'].values[0], row['ntrab_rl_otros'].values[0], row['ntrab_rl_noprec'].values[0]
    suma_A = A0 + A1 + A2 + A3 + A4 + A5
    B0, B1, B2, B3, B4, B5 = [round(round(val / suma_A, 4) * 100,2) for val in (A0, A1, A2, A3, A4, A5)]

    
    
    # Variables C,D
    C0, C1, C2, C3, C4, C5, C6 = row['ntrab_tc_indet'].values[0], row['ntrab_tc_nattemp'].values[0], row['ntrab_tc_natacc'].values[0], row['ntrab_tc_obrserv'].values[0], row['ntrab_tc_tiempar'].values[0], row['ntrab_tc_otros'].values[0], row['ntrab_tc_noprec'].values[0]
    suma_C = C0 + C1 + C2 + C3 + C4 + C5 + C6
    D0, D1, D2, D3, D4, D5, D6 = [round(round(val / suma_C, 4) * 100,2) for val in (C0, C1, C2, C3, C4, C5, C6)]

    
    # Variables E,F
    E0,E1,E2,E3 = row['ntrab_rp_spp'].values[0], row['ntrab_rp_snp'].values[0], row['ntrab_rp_otros'].values[0], row['ntrab_rp_noprecrp'].values[0]
    suma_E = E0 + E1 + E2 + E3
    F0,F1,F2,F3 = [round(round(val / suma_E, 4) * 100,2) for val in (E0,E1,E2,E3)]

    # Variables G,H 
    G0,G1,G2,G3 = row['ntrab_afil_essalud'].values[0], row['ntrab_afil_essalud_agrar'].values[0], row['ntrab_afil_sis'].values[0], row['ntrab_afil_noprec'].values[0]
    suma_G = G0 + G1 + G2 + G3
    H0, H1, H2, H3 = [round(round(val / suma_G, 4) * 100,2) for val in (G0,G1,G2,G3)]

    
    # Variables I,J
    I0,I1,I2 = row['ntrab_18mas'].values[0], row['ntrab_menos18'].values[0], row['ntrab_edad_noprec'].values[0]
    suma_I = I0 + I1 + I2
    J0, J1, J2 = [round(round(val / suma_I, 4) * 100,2) for val in (I0,I1,I2)]
    
    
    # Variables K,L
    K0,K1,K2 = row['ntrab_mujeres'].values[0], row['ntrab_hombres'].values[0], row['ntrab_sexo_noprec'].values[0]
    suma_K = K0 + K1 + K2
    L0, L1, L2 = [round(round(val / suma_K, 4) * 100,2)for val in (K0, K1, K2)]

    # Variables M,N
    M0, M1 = row['ntrab_sind'].values[0], row['ntrab_nosind'].values[0]
    suma_M = M0 + M1
    N0, N1 = [round(round(val / suma_M, 4) * 100,2) for val in (M0, M1)]


    # Variables O,P
    O0, O1, O2 = row['ntrab_consctr'].values[0], row['ntrab_sinsctr'].values[0], row['ntrab_sexo_noprec'].values[0]
    suma_O = O0 + O1 + O2
    P0, P1, P2 = [round(round(val / suma_O, 4) * 100,2) for val in (O0, O1, O2)]
    
    # Diccionario de reemplazos
    reemplazos = {
        'Z0': str(Z0),
        'Z1': str(Z1),
        'Z2': str(Z2),
        'Z3': str(Z3),
        'Z4': str(Z4),
        'Z5': str(Z5),
        'Z6': str(Z6),
        'Z7': str(Z7),
        'A0': str(A0),
        'A1': str(A1),
        'A2': str(A2),
        'A3': str(A3),
        'A4': str(A4),
        'A5': str(A5),
        'B0': str(B0),
        'B1': str(B1), 
        'B2': str(B2),
        'B3': str(B3),
        'B4': str(B4),
        'B5': str(B5),
        'C0': str(C0),
        'C1': str(C1),
        'C2': str(C2),
        'C3': str(C3),
        'C4': str(C4),
        'C5': str(C5),
        'C6': str(C6),
        'D0': str(D0),
        'D1': str(D1),
        'D2': str(D2),
        'D3': str(D3),
        'D4': str(D4),
        'D5': str(D5),
        'D6': str(D6),
        'E0': str(E0),
        'E1': str(E1),
        'E2': str(E2),
        'E3': str(E3),
        'F0': str(F0),
        'F1': str(F1),
        'F2': str(F2),
        'F3': str(F3),
        'G0': str(G0),
        'G1': str(G1),
        'G2': str(G2),
        'G3': str(G3),
        'H0': str(H0),
        'H1': str(H1),
        'H2': str(H2),
        'H3': str(H3),
        'I0': str(I0),
        'I1': str(I1),
        'I2': str(I2),
        'J0': str(J0),
        'J1': str(J1),
        'J2': str(J2),
        'K0': str(K0),
        'K1': str(K1),
        'K2': str(K2),
        'L0': str(L0),
        'L1': str(L1),
        'L2': str(L2),
        'M0': str(M0),
        'M1': str(M1),
        'N0': str(N0),
        'N1': str(N1),
        'O0': str(O0),
        'O1': str(O1),
        'O2': str(O2),
        'P0': str(P0),
        'P1': str(P1),
        'P2': str(P2)
    }
    
    
    replace_text(doc, reemplazos)
    
    doc.save(f"planilla_Electrónica_{str(fechas['fecha'])}.docx")
    print('Documento Generado')

def apply_format_to_cell(cell, bold=False, font_size=None):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if bold:
                run.bold = True
            if font_size:
                run.font.size = Pt(font_size)

def apply_format_to_text(doc, text_to_format, bold=False, font_size=None):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if text_to_format in run.text:
                if bold:
                    run.bold = True
                if font_size:
                    run.font.size = Pt(font_size)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if text_to_format in run.text:
                            if bold:
                                run.bold = True
                            if font_size:
                                run.font.size = Pt(font_size)





# ------------------------------------------------------

def generate_planilla_without_obs(aniomes):
    fechas = {"fecha":aniomes}
    row = df[df['V_NUMPERIOD']==fechas["fecha"]]

    formatos = [
    ("Planilla Electrónica", {'bold': False, 'font_size': 18}),
    ("Empresa:", {'bold': True, 'font_size': 12}),
    ("Departamento", {'bold': True, 'font_size': 12}),
    ("Sector", {'bold': True, 'font_size': 12}),
    ("Años en el mercado", {'bold': True, 'font_size': 12}),
    ("Costo Salarial", {'bold': True, 'font_size': 12}),
    ("SST", {'bold': True, 'font_size': 12}),
    ("Nº de trabajadores", {'bold': True, 'font_size': 12}),
    ("Régimen Laboral", {'bold': True, 'font_size': 11}),
    ("Régimen Pensionario", {'bold': True, 'font_size': 11}),
    ("Afiliación Sindical", {'bold': True, 'font_size': 11}),
    ("Rango de Edades", {'bold': True, 'font_size': 11}),
    ("Tipo de Contrato", {'bold': True, 'font_size': 11}),
    ("Régimen de Salud", {'bold': True, 'font_size': 11}),
    ("Género", {'bold': True, 'font_size': 11}),
    ("Afiliación al SCTR", {'bold': True, 'font_size': 11}),
    ("N° Trabj.", {'bold': True, 'font_size': 11}),
    ("% Trabj.", {'bold': True, 'font_size': 11}),

    ('Multas Laborales', {'bold': True, 'font_size': 12}),
    ('Nº de Trabj. Afectados', {'bold': True, 'font_size': 11}),
    ('Monto', {'bold': True, 'font_size': 11}),
    ('Fecha Multa', {'bold': True, 'font_size': 11}),
    ]


    doc = Document('C:/Users/alexa/Dataton/Labora-Tech-2024/backend/generate_pdf_planilla/template.docx')
    process_text(row, fechas, doc)


    doc.add_paragraph("")

    titulo = doc.add_paragraph("Multas Laborales")

    df_tabla = pd.read_csv('C:/Users/alexa/Dataton/Labora-Tech-2024/backend/generate_pdf_planilla/data/data2.csv')


    df_tabla.loc[:,'V_NUMPERIOD'] = df_tabla['V_NUMPERIOD'].astype(str)
    df_tabla.loc[:,'V_NUMPERIOD'] = df_tabla['V_NUMPERIOD'].apply(lambda x: f"{x[:4]}-{x[4:]}")

    data_multas = df_tabla[['V_NUMPERIOD', 'TRAB_AFEC', 'MONTO_MULTA']]
    data_multas.columns = ["periodo","Nº de Trabj. Afectados","Monto"]
    data_multas = data_multas.sort_values('periodo').set_index('periodo').T
    data_multas = data_multas.reset_index()
    data_multas.rename(columns={'index':'Fecha Multa'},inplace=True)

    tabla_multas = doc.add_table(rows=1 + len(data_multas), cols=len(data_multas.columns))
    tabla_multas.style = 'Table Grid'

    for j, col_name in enumerate(data_multas.columns):
        cell = tabla_multas.rows[0].cells[j]
        cell.text = str(col_name)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for i, row in data_multas.iterrows():
        for j, value in enumerate(row):
            cell = tabla_multas.rows[i + 1].cells[j]
            cell.text = str(value)  # Convertir todos los valores a cadena
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER




    for texto, propiedades in formatos:
        apply_format_to_text(doc, texto, **propiedades)

    doc.save(f"planilla_Electrónica_{str(fechas['fecha'])}.docx")
    nombre_docx = f"planilla_Electrónica_{str(fechas['fecha'])}.docx"
    convert(nombre_docx)

    # Nombre del archivo generado dinámicamente
    file_name = f"planilla_Electrónica_{str(fechas['fecha'])}.docx"

    # Verifica si el archivo existe antes de intentar eliminarlo
    if os.path.exists(file_name):
        os.remove(file_name)
        print(f"Archivo eliminado: {file_name}")
    else:
        print(f"El archivo no existe: {file_name}")